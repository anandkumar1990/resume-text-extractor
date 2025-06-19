from flask import jsonify
import os
import tempfile
from io import BytesIO
import fitz  
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import docx
import re
import subprocess
from tqdm import tqdm
import pdfplumber
from pdf2image import convert_from_bytes, convert_from_path
import cv2
import numpy as np
import string
import zipfile
import docx2txt

def sanitize_text(text):
    return " ".join(text.split())
def ocr_image(image: Image.Image) -> str:
    gray = image.convert("L")
    return pytesseract.image_to_string(gray)

def remove_control_chars(text):
    return ''.join(ch for ch in text if ch in string.printable or ch == '\n')

def filter_unwanted_lines(text):
    lines = text.splitlines()
    filtered = []
    for line in lines:
        l = line.strip()
        if not l:
            continue
        if l.lower().startswith("resume - styled html") or l.startswith("https://"):
            continue
        if l.lower().startswith("page ") or l.lower().startswith("--- page"):
            continue
        if sum(c.isalnum() for c in l) < 0.4 * len(l):
            continue
        # Remove PyMuPDF image placeholder lines like <image: DeviceRGB, width: 441, height: 260, bpc: 8>
        if l.startswith("<image:") and l.endswith(">"):
            continue
        filtered.append(l)
    return '\n'.join(filtered)

def convert_doc_to_docx(doc_path):
    try:
        output_dir = os.path.dirname(doc_path)
        subprocess.run(['soffice', '--headless', '--convert-to', 'docx', doc_path, '--outdir', output_dir],
                       stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        converted_path = os.path.splitext(doc_path)[0] + '.docx'
        return converted_path if os.path.exists(converted_path) else None
    except Exception as e:
        print(f"[ERROR] DOC to DOCX conversion failed: {e}")
        return None
def ocr_pdf(file_bytes: bytes) -> str:
    text = ""
    images = convert_from_bytes(file_bytes)
    for i, image in enumerate(images):
        page_text = ocr_image(image)
        text += f"\n--- Page {i+1} OCR Text ---\n{page_text}\n"
    return text

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            blocks = page.get_text("blocks")
            blocks_sorted = sorted(blocks, key=lambda b: (b[1], b[0]))
            page_text = ""
            for b in blocks_sorted:
                block_text = b[4].strip()
                if block_text:
                    page_text += block_text + "\n"

            # OCR on first page only if text is too short or looks invalid
            if page_num == 0:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = ocr_image(img)
                if len(page_text.strip()) < 50 or not any(x.isalpha() for x in page_text[:50]):
                    page_text = ocr_text  # Use OCR instead
                # else: keep existing page_text (no need to merge blindly)

            if page_text.strip():
                text += f"\n--- Page {page_num+1} ---\n{page_text}"
            else:
                # Use OCR if no extractable text found
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = ocr_image(img)
                text += f"\n--- Page {page_num+1} OCR Text ---\n{ocr_text}\n"

        text = remove_control_chars(text)
        text = filter_unwanted_lines(text)

    except Exception as e:
        print(f"[WARN] fitz failed: {e}")

    # Fallback: try pdfplumber if text is still empty
    if not text.strip():
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", mode="wb") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            os.unlink(tmp_path)
        except Exception as e:
            print(f"[WARN] pdfplumber failed: {e}")

    # Fallback: OCR all pages
    if not text.strip():
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", mode="wb") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            images = convert_from_path(tmp_path)
            for img in images:
                img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
                ocr_text = pytesseract.image_to_string(img_cv)
                text += ocr_text + "\n"
            print(f"[INFO] OCR used for PDF bytes")
            os.unlink(tmp_path)
        except Exception as e:
            print(f"[ERROR] OCR failed for PDF: {e}")

    return sanitize_text(text)
def extract_text_from_docx(file_bytes: bytes) -> str:
    import zipfile

    text = ""
    tmp_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode="wb") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        # First attempt: docx2txt
        try:
            text = docx2txt.process(tmp_path)
            if len(text.strip()) > 500:
                os.unlink(tmp_path)
                return sanitize_text(text)
        except Exception as e:
            print(f"[WARN] docx2txt failed: {e}")

        # If docx2txt failed or gave poor content, do manual extraction
        full_text_parts = []

        # 1. Extract headers and footers
        try:
            doc = docx.Document(tmp_path)
            headers_footers = []
            for section in doc.sections:
                if section.header:
                    headers_footers += [p.text.strip() for p in section.header.paragraphs if p.text.strip()]
                if section.footer:
                    headers_footers += [p.text.strip() for p in section.footer.paragraphs if p.text.strip()]
            if headers_footers:
                full_text_parts.append("\n".join(headers_footers))
        except Exception as e:
            print(f"[WARN] Could not extract headers/footers: {e}")

        # 2. OCR embedded images
        try:
            images_text = []
            with zipfile.ZipFile(tmp_path, 'r') as z:
                for file in z.namelist():
                    if file.startswith("word/media/") and file.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff")):
                        img_data = z.read(file)
                        img = Image.open(BytesIO(img_data))
                        ocr_txt = ocr_image(img)
                        if ocr_txt.strip():
                            images_text.append(f"[Embedded image OCR text start]\n{ocr_txt}\n[Embedded image OCR text end]")
            if images_text:
                full_text_parts.append("\n\n".join(images_text))
        except Exception as e:
            print(f"[WARN] Could not extract image OCR from DOCX: {e}")

        # 3. Fallback: docx paragraph text
        try:
            doc = docx.Document(tmp_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                full_text_parts.append("\n".join(paragraphs))
        except Exception as e:
            print(f"[ERROR] python-docx fallback failed: {e}")

        os.unlink(tmp_path)
        return sanitize_text("\n\n".join(full_text_parts))

    except Exception as e:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        print(f"[ERROR] extract_text_from_docx failed: {e}")
        return ""

def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"[ERROR] Failed to decode TXT file: {e}")
        return ""

def extract_embedded_image_text(docx_path):
    extracted_text = ""
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            for file in z.namelist():
                if file.startswith("word/media/") and file.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff")):
                    img_data = z.read(file)
                    img = Image.open(BytesIO(img_data))
                    extracted_text += ocr_image(img) + "\n"
    except Exception as e:
        print(f"[WARN] Failed to extract OCR from embedded images: {e}")
    return extracted_text

def clean_text(text: str) -> str:
    # Replace multiple newlines/tabs with space
    text = re.sub(r'[\n\r\t]+', ' ', text)
    # Remove multiple spaces
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()


# Use environment variables or fallback to a relative path
INPUT_DIR = os.environ.get("INPUT_DIR", "./input_resumes")
OUTPUT_DIR = os.environ.get("OUTPUT_DIR", "./output_resumes")

def process_resume(file_path, output_dir):
    ext = file_path.lower().split('.')[-1]
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    output_path = os.path.join(output_dir, base_name + ".txt")

    if os.path.exists(output_path):
        print(f"[SKIPPED] Already processed: {base_name}")
        return True, file_path, "Already processed"

    text = ""

    try:
        with open(file_path, "rb") as f:
            file_bytes = f.read()
    except Exception as e:
        msg = f"Failed to read file: {e}"
        print(f"[ERROR] {msg}")
        return False, file_path, msg

    if ext == "pdf":
        try:
            text = extract_text_from_pdf(file_bytes)
        except Exception as e:
            msg = f"extract_text_from_pdf failed: {e}"
            print(f"[WARN] {msg}")
            return False, file_path, msg
    elif ext == "docx":
        try:
            text = extract_text_from_docx(file_bytes)
        except Exception as e:
            msg = f"extract_text_from_docx failed: {e}"
            print(f"[WARN] {msg}")
            return False, file_path, msg
    elif ext == "doc":
        try:
            converted_path = convert_doc_to_docx(file_path)
            if converted_path:
                with open(converted_path, "rb") as f:
                    docx_bytes = f.read()
                text = extract_text_from_docx(docx_bytes)
            else:
                msg = f"DOC conversion failed"
                print(f"[ERROR] {msg}: {file_path}")
                return False, file_path, msg
        except Exception as e:
            msg = f"extract_text_from_doc fallback failed: {e}"
            print(f"[ERROR] {msg}")
            return False, file_path, msg
    elif ext == "txt":
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            msg = f"TXT read failed: {e}"
            print(f"[ERROR] {msg}")
            return False, file_path, msg
    else:
        msg = f"Unsupported file type: {ext}"
        print(f"[ERROR] {msg}")
        return False, file_path, msg

    if text and text.strip():
        sanitized = sanitize_text(text)
        try:
            with open(output_path, "w", encoding="utf-8", errors="replace") as out:
                out.write(sanitized)
        except Exception as e:
            msg = f"Failed to write output TXT: {e}"
            print(f"[ERROR] {msg} for {base_name}")
            return False, file_path, msg
        return True, file_path, "Success"
    else:
        msg = "No extractable content"
        print(f"[WARN] {msg} in: {file_path}")
        return False, file_path, msg

def extract_all_resumes(input_dir, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    failed_files = []
    failed_reasons = {}  # filename (base) -> reason
    resume_files = [f for f in os.listdir(input_dir) if f.lower().endswith(('.pdf', '.docx', '.doc', '.txt'))]

    for file in tqdm(resume_files, desc="Processing Resumes"):
        full_path = os.path.join(input_dir, file)
        success, name, reason = process_resume(full_path, output_dir)
        if not success:
            failed_files.append((name, reason))
            failed_reasons[os.path.splitext(os.path.basename(name))[0]] = reason

    if failed_files:
        print("\n⚠️ Failed or Incomplete Extraction for:")
        for name, reason in failed_files:
            print(f" - {os.path.basename(name)} | Reason: {reason}")
    else:
        print("\n✅ All resumes processed successfully.")

    return failed_reasons  # <-- return for later use
INPUT_DIR = r"C:\Users\AnandaKumarD\Desktop\CLEAN-RESUMES\input_resumes"
OUTPUT_DIR = r"C:\Users\AnandaKumarD\Desktop\CLEAN-RESUMES\output_resumes"

def batch_process_files(input_dir, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    processed = []
    skipped = []
    seen = set()
    resume_files = [f for f in os.listdir(input_dir) if f.lower().endswith(('.pdf', '.docx', '.doc', '.txt'))]
    for file in resume_files:
        key = file.lower()  # skip only exact name+ext duplicates
        base_name = os.path.splitext(os.path.basename(file))[0]
        output_txt_path = os.path.join(output_dir, base_name + ".txt")
        if key in seen or os.path.exists(output_txt_path):
            skipped.append(file)
            continue
        seen.add(key)
        full_path = os.path.join(input_dir, file)
        success, _, _ = process_resume(full_path, output_dir)
        if success:
            processed.append(file)
        else:
            skipped.append(file)
    print(f"Processed: {len(processed)}, Skipped: {len(skipped)}")  # <-- Add this line
    return {
        "processed_count": len(processed),
        "skipped_count": len(skipped),
        "skipped_files": skipped
    }

